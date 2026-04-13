from __future__ import annotations

import argparse
from copy import deepcopy
from dataclasses import dataclass
import importlib.util
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docxcompose.composer import Composer
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor


DEFAULT_TEMPLATE = Path(__file__).with_name("pandoc_docx_defaults.yaml")
DEFAULT_REFERENCE_DOC = Path(__file__).with_name("pandoc_reference.docx")
DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parents[1] / "pandoc_docx"

SECTION_ORDER = {
    "前言": 0,
    "自序": 1,
}

BODY_FONT_SIZE = 12
BODY_FIRST_LINE_INDENT = Pt(24)
BODY_LINE_SPACING = 1.5
TOP_BOTTOM_MARGIN_CM = 2.6
LEFT_MARGIN_CM = 2.8
RIGHT_MARGIN_CM = 2.4
HEADER_FOOTER_DISTANCE_CM = 1.5
MAX_IMAGE_PORTION = 3 / 4
TABLE_WIDTH_PORTION = 0.95
MIN_TABLE_COLUMN_WIDTH_CM = 1.8
DEFAULT_TABLE_OVERFLOW_THRESHOLD = 1.12
DEFAULT_TABLE_MIN_FONT_SIZE = 9.0
EQUATION_NUMBER_COLUMN_RATIO = 0.14
EQUATION_LAYOUT_MARK = "EquationLayout"

CHAPTER_TITLE_PATTERN = re.compile(r"^(前言|自序|第\s*[0-9一二三四五六七八九十]+\s*章)")
CAPTION_PREFIX_PATTERN = re.compile(
    r"^(图|表)\s*([0-9一二三四五六七八九十]+(?:[-—–.．][0-9一二三四五六七八九十]+)*)\s*[：:．。.]?\s*(.*)$"
)
EQUATION_TAG_PATTERN = re.compile(r"\\tag\*?\{([^{}]+)\}")
EQUATION_NUMBER_PATTERN = re.compile(r"^\([^)]+\)$")

CHINESE_DIGITS = {
    "零": 0,
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}


@dataclass(frozen=True)
class ImageSizingOptions:
    max_width_ratio: float
    max_height_ratio: float
    small_image_mode: str


@dataclass(frozen=True)
class TableSizingOptions:
    width_ratio: float
    overflow_threshold: float
    overflow_strategy: str
    min_font_size: float


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Convert markdown files under each subfolder to DOCX with Pandoc and merge them by chapter order."
    )
    parser.add_argument(
        "input_root",
        nargs="?",
        type=Path,
        help="Root directory containing chapter subfolders. If omitted, MD_BOOK_PATH from src/utils.py is used.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=DEFAULT_OUTPUT_DIR,
        help=f"Directory for converted files and merged output. Default: {DEFAULT_OUTPUT_DIR}",
    )
    parser.add_argument(
        "--output-name",
        default="book_complete.docx",
        help="Filename for the merged DOCX output.",
    )
    parser.add_argument(
        "--defaults",
        type=Path,
        default=DEFAULT_TEMPLATE,
        help="Pandoc defaults YAML template path.",
    )
    parser.add_argument(
        "--reference-doc",
        type=Path,
        default=DEFAULT_REFERENCE_DOC,
        help="Pandoc reference DOCX path. It will be created automatically when missing.",
    )
    parser.add_argument(
        "--keep-intermediate",
        action="store_true",
        help="Keep per-markdown converted DOCX files.",
    )
    parser.add_argument(
        "--image-max-width-page-ratio",
        type=float,
        default=MAX_IMAGE_PORTION,
        help="Maximum image width as a ratio of printable page width.",
    )
    parser.add_argument(
        "--image-max-height-page-ratio",
        type=float,
        default=MAX_IMAGE_PORTION,
        help="Maximum image height as a ratio of printable page height.",
    )
    parser.add_argument(
        "--image-small-mode",
        choices=["keep", "enlarge"],
        default="keep",
        help="Whether small images keep original size or scale up toward the configured limits.",
    )
    parser.add_argument(
        "--table-width-page-ratio",
        type=float,
        default=TABLE_WIDTH_PORTION,
        help="Target maximum table width as a ratio of printable page width.",
    )
    parser.add_argument(
        "--table-overflow-threshold",
        type=float,
        default=DEFAULT_TABLE_OVERFLOW_THRESHOLD,
        help="If the estimated natural table width exceeds this ratio of the allowed width, trigger overflow handling.",
    )
    parser.add_argument(
        "--table-overflow-strategy",
        choices=["shrink-font", "auto"],
        default="auto",
        help="How to handle tables whose estimated natural width exceeds the threshold.",
    )
    parser.add_argument(
        "--table-min-font-size",
        type=float,
        default=DEFAULT_TABLE_MIN_FONT_SIZE,
        help="Minimum table font size when overflow handling chooses font shrinking.",
    )
    return parser.parse_args()


def build_image_sizing_options(args: argparse.Namespace) -> ImageSizingOptions:
    return ImageSizingOptions(
        max_width_ratio=max(args.image_max_width_page_ratio, 0.1),
        max_height_ratio=max(args.image_max_height_page_ratio, 0.1),
        small_image_mode=args.image_small_mode,
    )


def build_table_sizing_options(args: argparse.Namespace) -> TableSizingOptions:
    strategy = args.table_overflow_strategy
    if strategy == "auto":
        # Keep page layout consistent: auto falls back to non-page-changing mode.
        strategy = "shrink-font"

    return TableSizingOptions(
        width_ratio=max(min(args.table_width_page_ratio, 1.0), 0.3),
        overflow_threshold=max(args.table_overflow_threshold, 1.0),
        overflow_strategy=strategy,
        min_font_size=max(args.table_min_font_size, 6.0),
    )


def load_md_book_path_from_utils() -> Path | None:
    repo_root = Path(__file__).resolve().parents[2]
    if str(repo_root) not in sys.path:
        sys.path.insert(0, str(repo_root))

    try:
        from src.utils import MD_BOOK_PATH
    except Exception as exc:
        utils_path = repo_root / "src" / "utils.py"
        try:
            spec = importlib.util.spec_from_file_location(
                "md_book_reviser_utils", utils_path
            )
            if spec is None or spec.loader is None:
                raise ImportError(f"Could not create import spec for {utils_path}")
            utils_module = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(utils_module)
            MD_BOOK_PATH = getattr(utils_module, "MD_BOOK_PATH", None)
        except Exception as fallback_exc:
            print(
                "[warn] Failed to load MD_BOOK_PATH from src/utils.py: "
                f"{exc}; fallback also failed: {fallback_exc}"
            )
            return None

    if not MD_BOOK_PATH:
        return None
    return Path(MD_BOOK_PATH).expanduser()


def rewrite_equation_tags_for_docx(content: str) -> str:
    def replace_tag(match: re.Match[str]) -> str:
        label = match.group(1).strip()
        if not label:
            return match.group(0)
        display_label = (
            label if (label.startswith("(") and label.endswith(")")) else f"({label})"
        )
        return rf"\qquad \text{{{display_label}}}"

    return EQUATION_TAG_PATTERN.sub(replace_tag, content)


def create_pandoc_input(markdown_path: Path) -> Path:
    original_content = markdown_path.read_text(encoding="utf-8")
    processed_content = rewrite_equation_tags_for_docx(original_content)
    if processed_content == original_content:
        return markdown_path

    with tempfile.NamedTemporaryFile(
        mode="w",
        encoding="utf-8",
        suffix=".md",
        prefix=f"{markdown_path.stem}_pandoc_",
        dir=str(markdown_path.parent),
        delete=False,
    ) as temp_file:
        temp_file.write(processed_content)
        return Path(temp_file.name)


def chinese_to_int(value: str) -> int | None:
    if not value:
        return None
    if value.isdigit():
        return int(value)
    if value == "十":
        return 10
    if "十" in value:
        left, _, right = value.partition("十")
        tens = CHINESE_DIGITS.get(left, 1) if left else 1
        ones = CHINESE_DIGITS.get(right, 0) if right else 0
        return tens * 10 + ones
    total = 0
    for char in value:
        digit = CHINESE_DIGITS.get(char)
        if digit is None:
            return None
        total = total * 10 + digit
    return total


def natural_sort_key(text: str) -> list[object]:
    parts = re.split(r"(\d+)", text)
    return [int(part) if part.isdigit() else part.lower() for part in parts]


def chapter_sort_key(path: Path) -> tuple[int, int, str]:
    name = path.name.strip()
    for label, index in SECTION_ORDER.items():
        if label in name:
            return index, 0, name

    match = re.search(r"第\s*([0-9一二三四五六七八九十]+)\s*章", name)
    if match:
        chapter_number = chinese_to_int(match.group(1))
        if chapter_number is not None:
            return 10, chapter_number, name

    return 99, 0, name


def set_run_fonts(
    run, latin_font: str = "Times New Roman", east_asia_font: str = "宋体"
) -> None:
    run.font.name = latin_font
    run_properties = run._element.get_or_add_rPr()
    run_properties.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.color.rgb = RGBColor(0, 0, 0)


def configure_style(
    document: Document,
    style_name: str,
    *,
    size: int,
    bold: bool,
    italic: bool,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    first_line_indent=None,
    line_spacing: float | None = None,
    line_spacing_rule: WD_LINE_SPACING | None = None,
    space_before: int | None = None,
    space_after: int | None = None,
) -> None:
    try:
        style = document.styles[style_name]
    except KeyError:
        return

    if style.type != WD_STYLE_TYPE.PARAGRAPH:
        return

    font = style.font
    font.name = "Times New Roman"
    style_properties = style.element.get_or_add_rPr()
    style_properties.rFonts.set(qn("w:eastAsia"), "宋体")
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = RGBColor(0, 0, 0)

    paragraph_format = style.paragraph_format
    if alignment is not None:
        paragraph_format.alignment = alignment
    if first_line_indent is not None:
        paragraph_format.first_line_indent = first_line_indent
    if line_spacing_rule is not None:
        paragraph_format.line_spacing_rule = line_spacing_rule
    if line_spacing is not None:
        paragraph_format.line_spacing = line_spacing
    if space_before is not None:
        paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph_format.space_after = Pt(space_after)


def configure_document_layout(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(TOP_BOTTOM_MARGIN_CM)
        section.bottom_margin = Cm(TOP_BOTTOM_MARGIN_CM)
        section.left_margin = Cm(LEFT_MARGIN_CM)
        section.right_margin = Cm(RIGHT_MARGIN_CM)
        section.header_distance = Cm(HEADER_FOOTER_DISTANCE_CM)
        section.footer_distance = Cm(HEADER_FOOTER_DISTANCE_CM)


def configure_document_styles(document: Document) -> None:
    configure_style(
        document,
        "Normal",
        size=BODY_FONT_SIZE,
        bold=False,
        italic=False,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=BODY_FIRST_LINE_INDENT,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=0,
        space_after=0,
    )
    configure_style(
        document,
        "Body Text",
        size=BODY_FONT_SIZE,
        bold=False,
        italic=False,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=BODY_FIRST_LINE_INDENT,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=0,
        space_after=0,
    )
    configure_style(
        document,
        "Title",
        size=20,
        bold=True,
        italic=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=24,
        space_after=18,
    )
    configure_style(
        document,
        "Subtitle",
        size=14,
        bold=False,
        italic=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=12,
        space_after=12,
    )
    configure_style(
        document,
        "Heading 1",
        size=16,
        bold=True,
        italic=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=24,
        space_after=18,
    )
    configure_style(
        document,
        "Heading 2",
        size=14,
        bold=True,
        italic=False,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=18,
        space_after=12,
    )
    configure_style(
        document,
        "Heading 3",
        size=13,
        bold=True,
        italic=False,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=12,
        space_after=8,
    )
    configure_style(
        document,
        "Heading 4",
        size=12,
        bold=True,
        italic=False,
        line_spacing=BODY_LINE_SPACING,
        line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
        space_before=10,
        space_after=6,
    )
    configure_style(
        document,
        "Caption",
        size=11,
        bold=False,
        italic=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.0,
        line_spacing_rule=WD_LINE_SPACING.SINGLE,
        space_before=6,
        space_after=6,
    )


def get_content_extent(section) -> tuple[int, int]:
    available_width = int(
        section.page_width - section.left_margin - section.right_margin
    )
    available_height = int(
        section.page_height - section.top_margin - section.bottom_margin
    )
    return available_width, available_height


def get_max_image_extent(
    document: Document, image_options: ImageSizingOptions
) -> tuple[int, int]:
    section = document.sections[0]
    available_width, available_height = get_content_extent(section)
    return (
        int(available_width * image_options.max_width_ratio),
        int(available_height * image_options.max_height_ratio),
    )


def scale_inline_shapes(document: Document, image_options: ImageSizingOptions) -> None:
    max_width, max_height = get_max_image_extent(document, image_options)
    for shape in document.inline_shapes:
        current_width = int(shape.width)
        current_height = int(shape.height)
        if not current_width or not current_height:
            continue

        width_ratio = max_width / current_width
        height_ratio = max_height / current_height
        if image_options.small_image_mode == "enlarge":
            scale_ratio = min(width_ratio, height_ratio)
        else:
            width_scale = width_ratio if current_width > max_width else 1.0
            height_scale = height_ratio if current_height > max_height else 1.0
            scale_ratio = min(width_scale, height_scale)

        if scale_ratio == 1.0:
            continue

        shape.width = Emu(int(current_width * scale_ratio))
        shape.height = Emu(int(current_height * scale_ratio))


def estimate_text_units(text: str) -> int:
    total_units = 0
    for char in text.strip():
        total_units += 2 if ord(char) > 127 else 1
    return max(total_units, 1)


def estimate_table_natural_width(column_units: list[int], font_size_pt: float) -> int:
    point_width_per_unit = Pt(font_size_pt * 0.8)
    min_column_width = Cm(MIN_TABLE_COLUMN_WIDTH_CM)
    widths = [
        max(int(point_width_per_unit * units), int(min_column_width))
        for units in column_units
    ]
    return sum(widths)


def compute_table_font_size(
    overflow_ratio: float, table_options: TableSizingOptions
) -> float | None:
    if overflow_ratio <= table_options.overflow_threshold:
        return None

    target_font_size = BODY_FONT_SIZE / overflow_ratio
    return max(table_options.min_font_size, round(target_font_size, 1))


def set_table_font_size(table, font_size: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_fonts(run)
                    run.font.size = Pt(font_size)
                    run.font.italic = False


def set_table_no_borders(table) -> None:
    table_properties = table._element.tblPr
    existing_borders = table_properties.find(qn("w:tblBorders"))
    if existing_borders is not None:
        table_properties.remove(existing_borders)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    table_properties.append(borders)


def set_table_borders_without_outer_sides(table) -> None:
    table_properties = table._element.tblPr
    existing_borders = table_properties.find(qn("w:tblBorders"))
    if existing_borders is not None:
        table_properties.remove(existing_borders)

    borders = OxmlElement("w:tblBorders")
    # Word border size uses eighths of a point: 1.5pt=12, 1pt=8, 0.5pt=4.
    default_inner_line_width = 4
    thick_line_width = 12

    def append_border(edge: str, value: str, size: int | None = None) -> None:
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), value)
        if value != "nil":
            border.set(
                qn("w:sz"), str(size if size is not None else default_inner_line_width)
            )
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
        borders.append(border)

    append_border("top", "single", thick_line_width)
    append_border("left", "nil")
    append_border("bottom", "single", thick_line_width)
    append_border("right", "nil")
    append_border("insideH", "single", default_inner_line_width)
    append_border("insideV", "single", default_inner_line_width)

    table_properties.append(borders)


def set_table_first_inner_horizontal_border(table) -> None:
    if len(table.rows) < 2:
        return

    first_inner_horizontal_width = 8
    for cell in table.rows[0].cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        existing_borders = cell_properties.find(qn("w:tcBorders"))
        if existing_borders is not None:
            cell_properties.remove(existing_borders)

        borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(first_inner_horizontal_width))
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)
        cell_properties.append(borders)


def mark_equation_layout_table(table) -> None:
    table_properties = table._element.tblPr
    existing_caption = table_properties.find(qn("w:tblCaption"))
    if existing_caption is not None:
        table_properties.remove(existing_caption)

    caption = OxmlElement("w:tblCaption")
    caption.set(qn("w:val"), EQUATION_LAYOUT_MARK)
    table_properties.append(caption)


def is_equation_layout_table(table) -> bool:
    table_properties = table._element.tblPr
    caption = table_properties.find(qn("w:tblCaption"))
    return caption is not None and caption.get(qn("w:val")) == EQUATION_LAYOUT_MARK


def extract_equation_number(
    paragraph_element,
) -> tuple[object, str] | tuple[None, None]:
    math_para = paragraph_element.find(qn("m:oMathPara"))
    if math_para is not None:
        math_node = math_para.find(qn("m:oMath"))
    else:
        math_node = paragraph_element.find(qn("m:oMath"))

    if math_node is None:
        return None, None

    runs = list(math_node.findall(qn("m:r")))
    if not runs:
        return None, None

    last_run = runs[-1]
    texts = list(last_run.findall(qn("m:t")))
    if len(texts) != 1:
        return None, None

    label = texts[0].text or ""
    if not EQUATION_NUMBER_PATTERN.match(label.strip()):
        return None, None

    math_node.remove(last_run)
    while True:
        trailing_runs = list(math_node.findall(qn("m:r")))
        if not trailing_runs:
            break
        candidate_run = trailing_runs[-1]
        candidate_texts = list(candidate_run.findall(qn("m:t")))
        if not candidate_texts:
            break
        candidate_value = "".join(text.text or "" for text in candidate_texts)
        if candidate_value.strip():
            break
        math_node.remove(candidate_run)

    return paragraph_element, label.strip()


def build_right_aligned_label_paragraph(label: str):
    paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")

    justification = OxmlElement("w:jc")
    justification.set(qn("w:val"), "right")
    paragraph_properties.append(justification)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.append(spacing)

    indent = OxmlElement("w:ind")
    indent.set(qn("w:firstLine"), "0")
    paragraph_properties.append(indent)
    paragraph.append(paragraph_properties)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    run_properties.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(BODY_FONT_SIZE * 2))
    run_properties.append(size)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    run_properties.append(color)
    run.append(run_properties)

    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    paragraph.append(run)
    return paragraph


def convert_equation_numbers_to_right_aligned_layout(
    document: Document, section
) -> None:
    available_width, _ = get_content_extent(section)
    number_width = int(available_width * EQUATION_NUMBER_COLUMN_RATIO)
    formula_width = max(available_width - number_width, int(available_width * 0.7))

    body = document._element.body
    for child in list(body.iterchildren()):
        if child.tag != qn("w:p"):
            continue

        cloned_paragraph = deepcopy(child)
        equation_paragraph, label = extract_equation_number(cloned_paragraph)
        if equation_paragraph is None or label is None:
            continue

        layout_table = document.add_table(rows=1, cols=2)
        layout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        layout_table.autofit = False
        mark_equation_layout_table(layout_table)
        set_table_no_borders(layout_table)

        left_cell = layout_table.cell(0, 0)
        right_cell = layout_table.cell(0, 1)
        left_cell.width = Emu(formula_width)
        right_cell.width = Emu(number_width)
        layout_table.columns[0].width = Emu(formula_width)
        layout_table.columns[1].width = Emu(number_width)

        left_cell._tc.clear_content()
        left_cell._tc.append(equation_paragraph)

        right_cell._tc.clear_content()
        right_cell._tc.append(build_right_aligned_label_paragraph(label))

        child.addprevious(layout_table._element)
        child.getparent().remove(child)


def adjust_table_layout(table, section, table_options: TableSizingOptions) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    column_count = len(table.columns)
    if column_count == 0:
        return

    column_units = [1] * column_count
    for row in table.rows:
        for index, cell in enumerate(row.cells[:column_count]):
            cell_text = " ".join(
                paragraph.text.strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )
            column_units[index] = max(
                column_units[index], estimate_text_units(cell_text)
            )

    min_column_width = Cm(MIN_TABLE_COLUMN_WIDTH_CM)
    available_width, _ = get_content_extent(section)
    max_table_width = int(available_width * table_options.width_ratio)
    natural_width = estimate_table_natural_width(column_units, BODY_FONT_SIZE)
    overflow_ratio = natural_width / max_table_width if max_table_width else 1.0

    font_size = None
    if overflow_ratio > table_options.overflow_threshold:
        if table_options.overflow_strategy == "shrink-font":
            font_size = compute_table_font_size(overflow_ratio, table_options)
    if font_size is not None:
        set_table_font_size(table, font_size)

    total_units = sum(column_units) or column_count
    widths = [
        max(int(max_table_width * units / total_units), int(min_column_width))
        for units in column_units
    ]

    current_total = sum(widths)
    if current_total > max_table_width:
        scale_ratio = max_table_width / current_total
        widths = [
            max(int(width * scale_ratio), int(min_column_width)) for width in widths
        ]

    for column, width in zip(table.columns, widths):
        column.width = Emu(width)

    for row in table.rows:
        for index, cell in enumerate(row.cells[:column_count]):
            cell.width = Emu(widths[index])


def normalize_caption_text(text: str) -> str:
    stripped_text = text.strip()
    match = CAPTION_PREFIX_PATTERN.match(stripped_text)
    if not match:
        return text

    prefix, serial, caption_body = match.groups()
    normalized_serial = serial.replace("—", "-").replace("–", "-").replace("．", ".")
    caption_body = caption_body.strip()
    if caption_body:
        return f"{prefix}{normalized_serial} {caption_body}"
    return f"{prefix}{normalized_serial}"


def get_paragraph_style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def is_chapter_title(paragraph) -> bool:
    style_name = get_paragraph_style_name(paragraph)
    if style_name not in {"Heading 1", "Title"}:
        return False
    return bool(CHAPTER_TITLE_PATTERN.match(paragraph.text.strip()))


def apply_body_paragraph_format(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_reference_entry_paragraph_format(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Pt(0)
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def apply_heading_paragraph_format(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph_format.keep_with_next = True
    if is_chapter_title(paragraph):
        paragraph_format.space_before = Pt(24)
        paragraph_format.space_after = Pt(18)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(8)


def apply_caption_paragraph_format(paragraph) -> None:
    normalized_text = normalize_caption_text(paragraph.text)
    if normalized_text != paragraph.text:
        paragraph.text = normalized_text

    paragraph_format = paragraph.paragraph_format
    paragraph_format.first_line_indent = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.line_spacing = 1.0
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def ensure_reference_doc(reference_doc: Path) -> None:
    reference_doc.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE

    configure_document_layout(document)
    configure_document_styles(document)

    document.save(reference_doc)


def iter_paragraphs(document: Document) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def is_heading(paragraph) -> bool:
    style_name = get_paragraph_style_name(paragraph)
    return style_name.startswith("Heading") or style_name == "Title"


def is_caption(paragraph) -> bool:
    text = paragraph.text.strip()
    style_name = get_paragraph_style_name(paragraph)
    return style_name == "Caption" or bool(
        re.match(r"^(图|表)\s*[0-9一二三四五六七八九十]+", text)
    )


def contains_drawing(paragraph) -> bool:
    for run in paragraph.runs:
        if run._element.xpath(".//w:drawing"):
            return True
    return False


def is_reference_section_heading(paragraph) -> bool:
    return paragraph.text.strip() in {"参考文献", "相关链接"}


def has_numbering(paragraph) -> bool:
    paragraph_properties = paragraph._element.find(qn("w:pPr"))
    if paragraph_properties is None:
        return False
    return paragraph_properties.find(qn("w:numPr")) is not None


def strip_paragraph_numbering(paragraph) -> None:
    paragraph_properties = paragraph._element.get_or_add_pPr()

    numbering = paragraph_properties.find(qn("w:numPr"))
    if numbering is not None:
        paragraph_properties.remove(numbering)

    tabs = paragraph_properties.find(qn("w:tabs"))
    if tabs is not None:
        paragraph_properties.remove(tabs)

    indent = paragraph_properties.find(qn("w:ind"))
    if indent is not None:
        paragraph_properties.remove(indent)


def strip_leading_whitespace_runs(paragraph) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue

        stripped_text = run.text.lstrip(" \t")
        if stripped_text != run.text:
            run.text = stripped_text

        if run.text:
            break


def prepend_plain_number(paragraph, index: int) -> None:
    strip_paragraph_numbering(paragraph)
    apply_reference_entry_paragraph_format(paragraph)
    strip_leading_whitespace_runs(paragraph)

    number_run = paragraph.add_run()
    number_run.text = f"{index}."
    set_run_fonts(number_run)
    number_run.font.size = Pt(BODY_FONT_SIZE)
    number_run.font.italic = False
    paragraph._p.insert(0, number_run._r)


def normalize_reference_section_lists(document: Document) -> None:
    in_reference_section = False
    entry_index = 0

    for paragraph in document.paragraphs:
        if is_reference_section_heading(paragraph):
            in_reference_section = True
            entry_index = 0
            continue

        if in_reference_section and is_heading(paragraph):
            in_reference_section = False
            entry_index = 0

        if not in_reference_section or not has_numbering(paragraph):
            continue

        entry_index += 1
        prepend_plain_number(paragraph, entry_index)


def is_reference_section_entry(paragraph, in_reference_section: bool) -> bool:
    return in_reference_section and bool(paragraph.text.strip())


def _make_page_number_footer(
    section,
    start_page: int | None = None,
    fmt: str = "decimal",
) -> None:
    """Add a centered PAGE number field to the section's default footer.

    Args:
        section: python-docx Section object.
        start_page: If given, reset the page counter to this value.
        fmt: OOXML page-number format string, e.g. "decimal" (1,2,3…) or
             "upperRoman" (I,II,III…).
    """
    footer = section.footer
    for para in footer.paragraphs:
        para.clear()

    paragraphs = footer.paragraphs
    if paragraphs:
        para = paragraphs[0]
    else:
        para = footer.add_paragraph()

    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    run_begin = para.add_run()
    set_run_fonts(run_begin)
    run_begin.font.size = Pt(11)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = para.add_run()
    set_run_fonts(run_instr)
    run_instr.font.size = Pt(11)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr._r.append(instr)

    run_end = para.add_run()
    set_run_fonts(run_end)
    run_end.font.size = Pt(11)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)

    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if start_page is not None or fmt != "decimal":
        if pg_num_type is None:
            pg_num_type = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num_type)
        if start_page is not None:
            pg_num_type.set(qn("w:start"), str(start_page))
        else:
            pg_num_type.attrib.pop(qn("w:start"), None)
        pg_num_type.set(qn("w:fmt"), fmt)
    else:
        if pg_num_type is not None:
            pg_num_type.attrib.pop(qn("w:start"), None)
            pg_num_type.attrib.pop(qn("w:fmt"), None)
            if not pg_num_type.attrib:
                sect_pr.remove(pg_num_type)


def add_page_numbers_from_chapter(document: Document) -> None:
    """
    Add centered bottom page numbers to all sections:
    - Sections before the first main chapter (前言, 自序, etc.) use lowercase
      Roman numerals (i, ii, iii…) starting from i.
    - The first main chapter section (第X章) resets to Arabic numeral 1.
    - Subsequent chapter sections continue the Arabic numbering.
    """
    body = document._element.body
    section_has_main_chapter: list[bool] = []
    current_has_main_chapter = False

    for element in body:
        if element.tag == qn("w:p"):
            pPr = element.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None:
                    style_val = pStyle.get(qn("w:val"), "")
                    if style_val in ("Heading1", "Title"):
                        text = "".join(t.text or "" for t in element.iter(qn("w:t")))
                        if re.search(
                            r"第\s*[0-9一二三四五六七八九十]+\s*章", text.strip()
                        ):
                            current_has_main_chapter = True
                sectPr_in_p = pPr.find(qn("w:sectPr")) if pPr is not None else None
                if sectPr_in_p is not None:
                    section_has_main_chapter.append(current_has_main_chapter)
                    current_has_main_chapter = False
        elif element.tag == qn("w:sectPr"):
            section_has_main_chapter.append(current_has_main_chapter)

    if not section_has_main_chapter:
        section_has_main_chapter = [current_has_main_chapter]

    first_chapter_idx = next(
        (i for i, has_chapter in enumerate(section_has_main_chapter) if has_chapter),
        None,
    )
    for idx, section in enumerate(document.sections):
        if first_chapter_idx is None or idx < first_chapter_idx:
            # Preface / front-matter sections: Roman numerals from i.
            _make_page_number_footer(
                section,
                start_page=(1 if idx == 0 else None),
                fmt="upperRoman",
            )
        elif idx == first_chapter_idx:
            # First body chapter: reset to Arabic 1.
            _make_page_number_footer(section, start_page=1, fmt="decimal")
        else:
            # Subsequent chapters: continue Arabic numbering.
            _make_page_number_footer(section, fmt="decimal")


def postprocess_docx(
    docx_path: Path,
    image_options: ImageSizingOptions,
    table_options: TableSizingOptions,
) -> None:
    document = Document(docx_path)

    configure_document_layout(document)
    configure_document_styles(document)
    normalize_reference_section_lists(document)
    scale_inline_shapes(document, image_options)
    base_section = document.sections[0]
    body_paragraphs = tuple(document.paragraphs)
    body_paragraph_elements = {paragraph._element for paragraph in body_paragraphs}

    in_reference_section = False

    for paragraph in body_paragraphs:
        if not paragraph.text.strip() and not contains_drawing(paragraph):
            continue

        if is_reference_section_heading(paragraph):
            in_reference_section = True
            apply_heading_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_run_fonts(run)
                run.font.italic = False
            continue

        if is_heading(paragraph):
            in_reference_section = False
            apply_heading_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_run_fonts(run)
                run.font.italic = False
            continue

        if is_caption(paragraph):
            apply_caption_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_run_fonts(run)
                run.font.italic = False
            continue

        if contains_drawing(paragraph):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_fonts(run)
            continue

        if is_reference_section_entry(paragraph, in_reference_section):
            apply_reference_entry_paragraph_format(paragraph)
            strip_leading_whitespace_runs(paragraph)
            for run in paragraph.runs:
                set_run_fonts(run)
                run.font.italic = False
            continue

        apply_body_paragraph_format(paragraph)

        for run in paragraph.runs:
            set_run_fonts(run)
            run.font.italic = False

    for paragraph in iter_paragraphs(document):
        if paragraph._element in body_paragraph_elements:
            continue

        if not paragraph.text.strip() and not contains_drawing(paragraph):
            continue

        if is_heading(paragraph):
            apply_heading_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_run_fonts(run)
                run.font.italic = False
            continue

        if is_caption(paragraph):
            apply_caption_paragraph_format(paragraph)
            for run in paragraph.runs:
                set_run_fonts(run)
                run.font.italic = False
            continue

        if contains_drawing(paragraph):
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_fonts(run)
            continue

        apply_body_paragraph_format(paragraph)

        for run in paragraph.runs:
            set_run_fonts(run)
            run.font.italic = False

    convert_equation_numbers_to_right_aligned_layout(document, base_section)

    for table in document.tables:
        if is_equation_layout_table(table):
            set_table_no_borders(table)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            right_cell = table.cell(0, 1)
            for paragraph in right_cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    set_run_fonts(run)
                    run.font.size = Pt(BODY_FONT_SIZE)
                    run.font.italic = False
            continue

        adjust_table_layout(table, base_section, table_options)
        set_table_borders_without_outer_sides(table)
        set_table_first_inner_horizontal_border(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing_rule = (
                        WD_LINE_SPACING.SINGLE
                    )
                    paragraph.paragraph_format.line_spacing = 1.0
                    if row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        set_run_fonts(run)
                        run.font.italic = False

    add_page_numbers_from_chapter(document)
    document.save(docx_path)


def collect_markdown_files(input_root: Path) -> list[Path]:
    chapter_dirs = [path for path in input_root.iterdir() if path.is_dir()]
    ordered_dirs = sorted(chapter_dirs, key=chapter_sort_key)

    markdown_files: list[Path] = []
    for chapter_dir in ordered_dirs:
        markdown_files.extend(
            sorted(
                chapter_dir.rglob("*.md"),
                key=lambda path: (
                    natural_sort_key(str(path.relative_to(chapter_dir))),
                    path.name.lower(),
                ),
            )
        )

    return markdown_files


def run_pandoc(
    markdown_path: Path, output_path: Path, defaults_path: Path, reference_doc: Path
) -> None:
    pandoc_executable = shutil.which("pandoc")
    if not pandoc_executable:
        raise FileNotFoundError("Pandoc is not installed or not available in PATH.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    pandoc_input = create_pandoc_input(markdown_path)

    try:
        command = [
            pandoc_executable,
            pandoc_input.name,
            "--defaults",
            str(defaults_path.resolve()),
            "--reference-doc",
            str(reference_doc.resolve()),
            "-o",
            str(output_path.resolve()),
        ]
        subprocess.run(command, cwd=markdown_path.parent, check=True)
    finally:
        if pandoc_input != markdown_path:
            pandoc_input.unlink(missing_ok=True)


def merge_docx_files(
    docx_files: list[tuple[str, Path]],
    output_path: Path,
    image_options: ImageSizingOptions,
    table_options: TableSizingOptions,
) -> None:
    if not docx_files:
        raise ValueError("No DOCX files were generated; nothing to merge.")

    first_chapter_key, first_docx = docx_files[0]
    master_document = Document(first_docx)
    composer = Composer(master_document)
    current_chapter_key = first_chapter_key
    for chapter_key, docx_file in docx_files[1:]:
        if chapter_key != current_chapter_key:
            current_chapter_key = chapter_key
            # Insert a section break at chapter boundary so next chapter starts on a new page.
            master_document.add_section(WD_SECTION_START.NEW_PAGE)
        composer.append(Document(docx_file))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))
    postprocess_docx(output_path, image_options, table_options)


def build_output_name(markdown_path: Path, input_root: Path) -> str:
    relative_path = markdown_path.relative_to(input_root)
    flattened_parts = [part for part in relative_path.parts[:-1] if part]
    flattened_parts.append(markdown_path.stem)
    safe_name = "__".join(flattened_parts)
    safe_name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", safe_name)
    return f"{safe_name}.docx"


def main() -> int:
    args = parse_args()
    image_options = build_image_sizing_options(args)
    table_options = build_table_sizing_options(args)
    input_root_path = args.input_root
    using_md_book_path_test = False
    if input_root_path is None:
        input_root_path = load_md_book_path_from_utils()
        if input_root_path is None:
            raise FileNotFoundError(
                "No input_root provided and failed to read MD_BOOK_PATH from src/utils.py"
            )
        using_md_book_path_test = True
        print(f"[test] Using MD_BOOK_PATH from src/utils.py: {input_root_path}")

    input_root = input_root_path.resolve()
    defaults_path = args.defaults.resolve()
    reference_doc = args.reference_doc.resolve()
    if using_md_book_path_test and args.output_dir == DEFAULT_OUTPUT_DIR:
        output_dir = input_root
        print(f"[test] Output directory defaults to MD_BOOK_PATH: {output_dir}")
    else:
        output_dir = args.output_dir.resolve()
    merged_output = output_dir / args.output_name
    intermediate_dir = output_dir / "intermediate"

    if not input_root.exists() or not input_root.is_dir():
        raise FileNotFoundError(
            f"Input root does not exist or is not a directory: {input_root}"
        )
    if not defaults_path.exists():
        raise FileNotFoundError(
            f"Pandoc defaults template does not exist: {defaults_path}"
        )

    ensure_reference_doc(reference_doc)

    markdown_files = collect_markdown_files(input_root)
    if not markdown_files:
        raise FileNotFoundError(
            f"No markdown files found under child folders of: {input_root}"
        )

    generated_docx_files: list[tuple[str, Path]] = []
    for markdown_path in markdown_files:
        chapter_key = markdown_path.relative_to(input_root).parts[0]
        output_path = intermediate_dir / build_output_name(markdown_path, input_root)
        print(f"[pandoc] {markdown_path} -> {output_path}")
        run_pandoc(markdown_path, output_path, defaults_path, reference_doc)
        postprocess_docx(output_path, image_options, table_options)
        generated_docx_files.append((chapter_key, output_path))

    merge_docx_files(generated_docx_files, merged_output, image_options, table_options)
    print(f"[merged] {merged_output}")

    if not args.keep_intermediate:
        for _, docx_file in generated_docx_files:
            docx_file.unlink(missing_ok=True)
        if intermediate_dir.exists():
            try:
                shutil.rmtree(intermediate_dir)
            except OSError as exc:
                print(
                    f"[warn] Failed to remove intermediate directory {intermediate_dir}: {exc}"
                )

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
