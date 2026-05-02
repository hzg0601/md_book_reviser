import unittest
from pathlib import Path
import tempfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import parse_xml

from src.md2docx.build_book_docx import (
    ImageSizingOptions,
    TableSizingOptions,
    apply_code_listing_paragraph_style,
    get_math_text,
    is_caption,
    is_algorithm_math_paragraph,
    is_algorithm_listing_table,
    looks_like_caption_text,
    postprocess_docx,
    preprocess_markdown_for_docx,
)


class BuildBookDocxTests(unittest.TestCase):
    def test_preprocess_algorithm_flow_block(self) -> None:
        source = r"""
$$
\begin{array}{l}
\textbf{算法2-1 Flash Attention V1的前向传播算法流程} \\
\hline
\textbf{条件: } \text{矩阵 } \mathbf{Q}, \mathbf{K}, \mathbf{V} \in \mathbb{R}^{N \times d}
\end{array}
$$
""".strip()

        processed = preprocess_markdown_for_docx(source)

        self.assertEqual(processed, source)

    def test_inline_algorithm_reference_is_not_caption(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("算法3-4相比于算法3-2的差异如下。")

        self.assertFalse(is_caption(paragraph, tuple(document.paragraphs), 0))

    def test_algorithm_title_before_listing_is_caption(self) -> None:
        document = Document()
        document.styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
        title_paragraph = document.add_paragraph("算法3-4 Flash Attention 流程")
        document.add_paragraph("body", style="Source Code")

        self.assertTrue(is_caption(title_paragraph, tuple(document.paragraphs), 0))

    def test_figure_reference_text_is_not_caption_like(self) -> None:
        self.assertFalse(
            looks_like_caption_text(
                "图2-12 中为一个2阶段Warpgroup的流水线示意图，同一颜色的色块为同一个阶段。"
            )
        )

    def test_table_caption_text_is_caption_like(self) -> None:
        self.assertTrue(looks_like_caption_text("表2-1 FlashAttention 参数设置"))

    def test_source_code_is_mapped_to_listing_style(self) -> None:
        document = Document()
        document.styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
        document.styles.add_style("代码清单", WD_STYLE_TYPE.PARAGRAPH)
        paragraph = document.add_paragraph("print('hello')")
        paragraph.style = document.styles["Source Code"]

        apply_code_listing_paragraph_style(document, paragraph)

        self.assertEqual(paragraph.style.name, "代码清单")

    def test_algorithm_math_paragraph_is_detected(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        paragraph._element.append(
            parse_xml(
                '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                '<m:oMath><m:r><m:t>算法2-1 Flash Attention 流程</m:t></m:r></m:oMath>'
                '</m:oMathPara>'
            )
        )

        self.assertEqual(get_math_text(paragraph), "算法2-1 Flash Attention 流程")
        self.assertTrue(is_algorithm_math_paragraph(paragraph))

    def test_algorithm_math_paragraph_text_extraction(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        paragraph._element.append(
            parse_xml(
                '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                '<m:oMath><m:r><m:t>普通公式</m:t></m:r></m:oMath>'
                '</m:oMathPara>'
            )
        )

        self.assertEqual(get_math_text(paragraph), "普通公式")
        self.assertFalse(is_algorithm_math_paragraph(paragraph))

    def test_algorithm_math_block_is_wrapped_into_listing_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "algorithm.docx"
            document = Document()
            document.styles.add_style("代码清单", WD_STYLE_TYPE.PARAGRAPH)
            paragraph = document.add_paragraph()
            paragraph._element.append(
                parse_xml(
                    '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                    '<m:oMath><m:r><m:t>算法2-1 Flash Attention 流程</m:t></m:r></m:oMath>'
                    '</m:oMathPara>'
                )
            )
            document.save(docx_path)

            postprocess_docx(
                docx_path,
                ImageSizingOptions(
                    max_width_ratio=0.75,
                    max_height_ratio=0.75,
                    small_image_mode="keep",
                ),
                TableSizingOptions(
                    width_ratio=0.95,
                    overflow_threshold=1.12,
                    overflow_strategy="shrink-font",
                    min_font_size=9.0,
                ),
            )

            output_document = Document(docx_path)
            self.assertEqual(len(output_document.paragraphs), 0)
            self.assertEqual(len(output_document.tables), 1)
            self.assertTrue(is_algorithm_listing_table(output_document.tables[0]))
            self.assertEqual(
                output_document.tables[0].cell(0, 0).paragraphs[0].style.name,
                "代码清单",
            )
            self.assertEqual(
                output_document.tables[0].cell(0, 0).paragraphs[0].paragraph_format.line_spacing_rule,
                WD_LINE_SPACING.SINGLE,
            )


if __name__ == "__main__":
    unittest.main()